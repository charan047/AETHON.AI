from auth.dependencies import get_current_user, require_editor
from auth.org_context import OrgContext, check_plan_limit, get_org_context
from database.models import User
from fastapi import Depends

from fastapi import APIRouter, HTTPException, BackgroundTasks, Request, Response, Query
from fastapi.responses import StreamingResponse
from io import BytesIO
import asyncio
import base64
import html
import logging
import re
import secrets
import socket
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from pydantic import BaseModel, Field
from typing import Optional, List, Any, Literal
from datetime import datetime
from uuid import uuid4
from urllib.parse import urlparse
from database import get_db
from database.db import AsyncSessionLocal
from database.models import (
    Agent,
    AgentMemoryEntry,
    Client,
    ExecutionStatus,
    IntegrationType,
    UserIntegration,
    Workflow,
    Execution,
    Message,
)
from middleware.rate_limit import limiter
from runtime.graph_builder import WorkflowExecutionStopped
from runtime.workflow_engine import WorkflowEngine
from services.websocket_manager import ws_manager
from services.integration_crypto import decrypt_config
from tools.communication.utils import get_gmail_service

router = APIRouter(dependencies=[Depends(get_current_user), Depends(get_org_context)])
logger = logging.getLogger(__name__)


def _is_celery_broker_reachable(broker_url: str | None) -> bool:
    if not broker_url:
        return False
    parsed = urlparse(broker_url)
    if parsed.scheme not in {"redis", "rediss"}:
        return True
    host = parsed.hostname or "localhost"
    port = parsed.port or 6379
    try:
        with socket.create_connection((host, port), timeout=0.25):
            return True
    except OSError:
        return False


class ExecutionCreate(BaseModel):
    input_message: str
    trigger: str = "manual"
    max_runtime_seconds: Optional[int] = Field(default=None, ge=1, le=86400)


class WorkflowRunRequest(BaseModel):
    input_message: str = ""
    input_values: dict[str, str] = Field(default_factory=dict)
    client_id: str | None = None
    trigger: str = "manual"
    max_runtime_seconds: Optional[int] = Field(default=None, ge=1, le=86400)


class MessageResponse(BaseModel):
    id: str
    execution_id: str
    from_agent: str
    to_agent: Optional[str]
    content: str
    role: str
    token_count: int
    timestamp: datetime
    msg_metadata: Any

    model_config = {"from_attributes": True}


class ExecutionResponse(BaseModel):
    id: str
    workflow_id: str
    client_id: Optional[str]
    client_name: Optional[str] = None
    parent_execution_id: Optional[str] = None
    trigger: str
    status: str
    input_message: str
    output: Optional[str] = None
    output_message: Optional[str]
    revision_number: int = 1
    ceo_feedback: Optional[str] = None
    started_at: datetime
    completed_at: Optional[datetime]
    approved_by: Optional[str] = None
    approved_at: Optional[datetime] = None
    approval_note: Optional[str] = None
    delivered_at: Optional[datetime] = None
    delivery_method: Optional[str] = None
    delivery_target: Optional[str] = None
    token_count: int
    cost: float
    error: Optional[str]
    warning: Optional[str] = None
    max_runtime_seconds: int

    model_config = {"from_attributes": True}


class ExecutionStepResponse(BaseModel):
    id: str
    execution_id: str
    org_id: str
    step_type: str
    content: str
    tool_name: Optional[str]
    tool_input: Any
    tool_output: Any
    tool_success: Optional[bool]
    step_index: int
    duration_ms: Optional[int]
    tokens_used: Optional[int]
    created_at: datetime
    timestamp: Optional[datetime] = None
    agent_id: Optional[str] = None
    agent_name: Optional[str] = None

    model_config = {"from_attributes": True}


class ExecutionDetailResponse(ExecutionResponse):
    workflow_name: Optional[str] = None
    agent_name: Optional[str] = None
    model_name: Optional[str] = None
    input: str
    steps: List[ExecutionStepResponse] = Field(default_factory=list)


class ExecutionRunResponse(BaseModel):
    id: str
    execution_id: str
    status: str
    websocket_channel: str
    message: str


class ExecutionRegenerateResponse(BaseModel):
    revision_id: str
    revision_number: int
    status: str


class ExecutionRevisionResponse(BaseModel):
    id: str
    revision_number: int
    status: str
    ceo_feedback: Optional[str] = None
    output: Optional[str] = None
    started_at: datetime
    approved_at: Optional[datetime] = None
    approved_by: Optional[str] = None


class ExecutionApproveRequest(BaseModel):
    note: Optional[str] = None


class ExecutionRegenerateRequest(BaseModel):
    feedback: str = Field(..., min_length=1, max_length=5000)


class ExecutionDeliverRequest(BaseModel):
    method: Literal["email", "google_doc", "portal"]
    email_to: Optional[str] = Field(default=None, max_length=255)
    doc_title: Optional[str] = Field(default=None, max_length=255)


class ExecutionDeliverResponse(BaseModel):
    delivered: bool
    method: str
    target: str
    delivered_at: datetime


def _strip_markdown_inline(text: str) -> str:
    value = text or ""
    value = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"[*_~#>]", "", value)
    return value.strip()


def _parse_markdown_blocks(content: str) -> list[dict[str, Any]]:
    lines = (content or "").replace("\r\n", "\n").split("\n")
    blocks: list[dict[str, Any]] = []
    paragraph: list[str] = []
    bullets: list[str] = []
    numbered: list[str] = []
    code_lines: list[str] = []
    in_code = False

    def flush_paragraph():
        nonlocal paragraph
        if paragraph:
            blocks.append({"type": "paragraph", "text": _strip_markdown_inline(" ".join(paragraph))})
            paragraph = []

    def flush_bullets():
        nonlocal bullets
        if bullets:
            blocks.append({"type": "bullet", "items": [_strip_markdown_inline(item) for item in bullets]})
            bullets = []

    def flush_numbered():
        nonlocal numbered
        if numbered:
            blocks.append({"type": "numbered", "items": [_strip_markdown_inline(item) for item in numbered]})
            numbered = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            code_lines = []

    def flush_text():
        flush_paragraph()
        flush_bullets()
        flush_numbered()

    for raw_line in lines:
        line = raw_line.rstrip()
        trimmed = line.strip()

        if trimmed.startswith("```"):
            flush_text()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(raw_line)
            continue

        if not trimmed:
            flush_text()
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", trimmed)
        if heading:
            flush_text()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading.group(1)),
                    "text": _strip_markdown_inline(heading.group(2)),
                }
            )
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", trimmed)
        if bullet:
            flush_paragraph()
            flush_numbered()
            bullets.append(bullet.group(1))
            continue

        numbered_match = re.match(r"^\d+\.\s+(.+)$", trimmed)
        if numbered_match:
            flush_paragraph()
            flush_bullets()
            numbered.append(numbered_match.group(1))
            continue

        flush_bullets()
        flush_numbered()
        paragraph.append(trimmed)

    flush_text()
    flush_code()
    return blocks


def _format_export_date(value: datetime | None) -> str:
    if not value:
        return datetime.utcnow().strftime("%B %d, %Y")
    return value.strftime("%B %d, %Y")


def _build_export_filename(workflow_name: str, completed_at: datetime | None, export_format: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", (workflow_name or "report").lower()).strip("-") or "report"
    date_part = (completed_at or datetime.utcnow()).strftime("%Y-%m-%d")
    return f"{slug}-{date_part}.{export_format}"


def _markdown_to_plain_text(content: str) -> str:
    lines: list[str] = []
    for block in _parse_markdown_blocks(content):
        if block["type"] == "heading":
            lines.append(block["text"])
        elif block["type"] in {"bullet", "numbered"}:
            prefix = "-" if block["type"] == "bullet" else "1."
            for item in block["items"]:
                lines.append(f"{prefix} {item}")
        else:
            lines.append(block["text"])
        lines.append("")
    return "\n".join(lines).strip()


def _markdown_to_html(content: str) -> str:
    parts: list[str] = []
    for block in _parse_markdown_blocks(content):
        if block["type"] == "heading":
            level = min(max(int(block["level"]), 1), 3)
            parts.append(f"<h{level}>{html.escape(block['text'])}</h{level}>")
        elif block["type"] == "bullet":
            items = "".join(f"<li>{html.escape(item)}</li>" for item in block["items"])
            parts.append(f"<ul>{items}</ul>")
        elif block["type"] == "numbered":
            items = "".join(f"<li>{html.escape(item)}</li>" for item in block["items"])
            parts.append(f"<ol>{items}</ol>")
        elif block["type"] == "code":
            parts.append(f"<pre><code>{html.escape(block['text'])}</code></pre>")
        else:
            parts.append(f"<p>{html.escape(block['text']).replace(chr(10), '<br/>')}</p>")
    return "\n".join(parts)


def _delivery_subject(workflow_name: str, client_name: str | None) -> str:
    if client_name:
        return f"[{workflow_name}] — Completed for {client_name}"
    return f"[{workflow_name}] — Completed"


async def _render_execution_pdf(
    *,
    agency_name: str,
    workflow_name: str,
    client_name: str | None,
    completed_at: datetime | None,
    agent_name: str,
    approved_by_name: str,
    approved_at: datetime | None,
    output: str,
) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:
        raise RuntimeError("PDF export dependency is not installed") from exc

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=1.0 * inch,
        bottomMargin=0.85 * inch,
        title=workflow_name,
        author=agency_name,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="AethonTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#0F172A"),
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AethonMeta",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#475569"),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AethonHeading2",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#0F172A"),
            spaceBefore=10,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AethonHeading3",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#1E293B"),
            spaceBefore=8,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AethonBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=16,
            textColor=colors.HexColor("#0F172A"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AethonCode",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#0F172A"),
            backColor=colors.HexColor("#F8FAFC"),
            borderColor=colors.HexColor("#CBD5E1"),
            borderWidth=0.5,
            borderPadding=8,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AethonCenteredMeta",
            parent=styles["AethonMeta"],
            alignment=TA_CENTER,
        )
    )

    title = workflow_name + (f" — {client_name}" if client_name else "")
    story = [
        Paragraph(title, styles["AethonTitle"]),
        Paragraph(_format_export_date(completed_at), styles["AethonMeta"]),
        Paragraph(f"Agent: {agent_name}", styles["AethonMeta"]),
        Paragraph(
            f"Approved by: {approved_by_name} at {approved_at.strftime('%b %d, %Y %I:%M %p') if approved_at else 'Unknown'}",
            styles["AethonMeta"],
        ),
        Spacer(1, 10),
    ]

    for block in _parse_markdown_blocks(output):
        if block["type"] == "heading":
            level = block["level"]
            style_name = "AethonHeading2" if level <= 2 else "AethonHeading3"
            story.append(Paragraph(block["text"], styles[style_name]))
        elif block["type"] == "bullet":
            items = [ListItem(Paragraph(item, styles["AethonBody"])) for item in block["items"]]
            story.append(ListFlowable(items, bulletType="bullet", leftIndent=18))
            story.append(Spacer(1, 8))
        elif block["type"] == "numbered":
            items = [ListItem(Paragraph(item, styles["AethonBody"])) for item in block["items"]]
            story.append(ListFlowable(items, bulletType="1", leftIndent=18))
            story.append(Spacer(1, 8))
        elif block["type"] == "code":
            story.append(Paragraph(block["text"].replace("\n", "<br/>"), styles["AethonCode"]))
        else:
            story.append(Paragraph(block["text"], styles["AethonBody"]))

    def draw_header_footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.HexColor("#475569"))
        canvas.drawString(doc.leftMargin, LETTER[1] - 0.55 * inch, agency_name)
        canvas.drawRightString(
            LETTER[0] - doc.rightMargin,
            0.5 * inch,
            f"Page {canvas.getPageNumber()}",
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_header_footer, onLaterPages=draw_header_footer)
    return buffer.getvalue()


async def _render_execution_docx(
    *,
    agency_name: str,
    workflow_name: str,
    client_name: str | None,
    completed_at: datetime | None,
    agent_name: str,
    approved_by_name: str,
    approved_at: datetime | None,
    output: str,
) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("DOCX export dependency is not installed") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title_run = title.add_run(workflow_name + (f" — {client_name}" if client_name else ""))
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.bold = True

    for text in (
        _format_export_date(completed_at),
        f"Agent: {agent_name}",
        f"Approved by: {approved_by_name} at {approved_at.strftime('%b %d, %Y %I:%M %p') if approved_at else 'Unknown'}",
    ):
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(9.5)

    document.add_paragraph()

    for block in _parse_markdown_blocks(output):
        if block["type"] == "heading":
            level = min(int(block["level"]), 3)
            p = document.add_paragraph(style=f"Heading {level}")
            run = p.add_run(block["text"])
            run.font.name = "Arial"
            run.bold = True
        elif block["type"] == "bullet":
            for item in block["items"]:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(item).font.name = "Arial"
        elif block["type"] == "numbered":
            for item in block["items"]:
                p = document.add_paragraph(style="List Number")
                p.add_run(item).font.name = "Arial"
        elif block["type"] == "code":
            p = document.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            p = document.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = agency_name
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def _send_execution_email_via_gmail(
    *,
    org_id: str,
    user_id: str,
    to_email: str,
    subject: str,
    html_body: str,
    text_body: str,
    pdf_bytes: bytes | None,
    pdf_filename: str,
    prefetched_config: dict | None = None,
) -> None:
    service = await get_gmail_service(org_id, user_id, prefetched_config=prefetched_config)
    message = MIMEMultipart("mixed")
    message["to"] = to_email
    message["subject"] = subject

    alternative = MIMEMultipart("alternative")
    alternative.attach(MIMEText(text_body, "plain", "utf-8"))
    alternative.attach(MIMEText(html_body, "html", "utf-8"))
    message.attach(alternative)

    if pdf_bytes:
        attachment = MIMEApplication(pdf_bytes, _subtype="pdf")
        attachment.add_header("Content-Disposition", "attachment", filename=pdf_filename)
        message.attach(attachment)

    raw = base64.urlsafe_b64encode(message.as_bytes()).decode()
    await asyncio.to_thread(
        lambda: service.users().messages().send(userId="me", body={"raw": raw}).execute()
    )


async def _create_google_doc(
    *,
    org_id: str,
    user_id: str,
    title: str,
    content: str,
    prefetched_config: dict | None = None,
) -> str:
    from tools.productivity.google_docs import GoogleDocsTool

    tool = GoogleDocsTool(config=prefetched_config or {})
    outcome = await tool.execute(
        {"title": title, "content": content},
        org_id=org_id,
        user_id=user_id,
    )
    if not outcome.success:
        raise ValueError(outcome.error or "Google Docs delivery failed")
    return str(outcome.result or "")


async def _extract_preferences(feedback: str, agent_id: str) -> str | None:
    """
    Extract reusable preferences from one-time feedback.
    Converts specific feedback into a general instruction sentence.
    """
    def _heuristic_extract(text: str) -> str | None:
        lowered = text.lower()
        instructions: list[str] = []

        under_words = re.search(r"(?:under|less than|max(?:imum)? of?)\s+(\d{2,4})\s+words?", lowered)
        if under_words:
            instructions.append(f"Keep responses under {under_words.group(1)} words.")
        elif any(token in lowered for token in ("shorter", "concise", "brief", "short and")):
            instructions.append("Keep responses concise.")

        if any(token in lowered for token in ("informal", "casual", "conversational", "friendly")):
            tone = "informal" if "informal" in lowered else "casual, conversational"
            instructions.append(f"Use {tone} tone.")
        elif "formal" in lowered and "too formal" in lowered:
            instructions.append("Avoid overly formal tone.")

        if not instructions:
            return None
        return " ".join(instructions)[:500]

    heuristic = _heuristic_extract(feedback)
    if heuristic:
        return heuristic

    try:
        from services.model_service import model_service

        llm = model_service._build_from_settings(temperature=0, max_tokens=150)
        prompt = (
            "Extract reusable writing preferences from this CEO feedback.\n"
            "Return a single instruction sentence the agent should always follow.\n"
            "If the feedback is too specific to one task (not generalizable), return 'none'.\n\n"
            f"Agent ID: {agent_id}\n"
            f"Feedback: {feedback}\n\n"
            "Preference (or 'none'):"
        )
        response = await llm.ainvoke(prompt)
        result = str(response.content or "").strip()
        if result.lower() == "none" or len(result) < 5:
            return None
        return result[:500]
    except Exception as exc:
        logger.warning("Preference extraction failed for agent %s: %s", agent_id, exc)
        return None


async def _store_ceo_preference(
    *,
    agent_id: str | None,
    org_id: str,
    preference_text: str | None,
    db: AsyncSession,
    source: str,
) -> None:
    if not agent_id or not preference_text:
        return

    existing = await db.scalar(
        select(AgentMemoryEntry).where(
            AgentMemoryEntry.agent_id == agent_id,
            AgentMemoryEntry.org_id == org_id,
            AgentMemoryEntry.memory_type == "ceo_preference",
            AgentMemoryEntry.always_inject == True,
            AgentMemoryEntry.content_preview == preference_text,
        )
    )
    if existing:
        return

    pref = AgentMemoryEntry(
        id=str(uuid4()),
        agent_id=agent_id,
        org_id=org_id,
        mem0_memory_id=f"local:{uuid4()}",
        content_preview=preference_text[:500],
        memory_type="ceo_preference",
        importance_score=1.0,
        always_inject=True,
        source=source,
        tags=["ceo_preference", source],
    )
    db.add(pref)
    await db.commit()


def _ordered_workflow_agents(workflow: Workflow | None) -> list[tuple[str, str]]:
    ordered: list[tuple[str, str]] = []
    seen: set[str] = set()
    for node in (workflow.nodes if workflow else []) or []:
        data = (node or {}).get("data", {}) or {}
        agent_id = data.get("agent_id")
        if not isinstance(agent_id, str) or not agent_id or agent_id in seen:
            continue
        seen.add(agent_id)
        label = data.get("label")
        ordered.append((agent_id, label if isinstance(label, str) and label.strip() else "Agent"))
    return ordered


def _primary_workflow_agent_id(workflow: Workflow | None) -> str | None:
    for node in (workflow.nodes if workflow else []) or []:
        data = (node or {}).get("data", {}) or {}
        agent_id = data.get("agent_id")
        if isinstance(agent_id, str) and agent_id:
            return agent_id
        agent_ids = data.get("agent_ids") or (node or {}).get("agent_ids") or []
        if agent_ids:
            first = agent_ids[0]
            if isinstance(first, str) and first:
                return first
    return None


def _execution_output_text(execution: Execution) -> str:
    if execution.output_message:
        return execution.output_message
    steps = sorted((execution.steps or []), key=lambda item: item.step_index, reverse=True)
    for step in steps:
        if step.step_type == "final_answer" and step.content:
            return step.content
    return ""


def _build_regeneration_prompt(execution: Execution, feedback: str) -> str:
    return (
        f"{execution.input_message}\n\n"
        "---\n"
        "CEO FEEDBACK ON PREVIOUS VERSION:\n"
        f"{feedback}\n\n"
        "Please revise your response addressing all feedback above. "
        "Keep what was good, fix what was flagged."
    )


async def _resolve_execution_root(execution: Execution, db: AsyncSession, org_id: str) -> Execution:
    current = execution
    seen: set[str] = set()
    while current.parent_execution_id and current.parent_execution_id not in seen:
        seen.add(current.id)
        parent = await db.scalar(
            select(Execution).where(
                Execution.id == current.parent_execution_id,
                Execution.org_id == org_id,
            )
        )
        if not parent:
            break
        current = parent
    return current


async def _load_execution_revision_chain(execution: Execution, db: AsyncSession, org_id: str) -> list[Execution]:
    root = await _resolve_execution_root(execution, db, org_id)
    candidate_rows = (
        await db.execute(
            select(Execution).where(
                Execution.org_id == org_id,
                Execution.workflow_id == root.workflow_id,
            )
        )
    ).scalars().all()

    by_parent: dict[str | None, list[Execution]] = {}
    for candidate in candidate_rows:
        by_parent.setdefault(candidate.parent_execution_id, []).append(candidate)

    ordered: list[Execution] = []
    stack = [root]
    seen: set[str] = set()
    while stack:
        current = stack.pop(0)
        if current.id in seen:
            continue
        seen.add(current.id)
        ordered.append(current)
        children = sorted(
            by_parent.get(current.id, []),
            key=lambda item: (item.revision_number or 1, item.started_at or datetime.min),
        )
        stack.extend(children)

    return sorted(ordered, key=lambda item: (item.revision_number or 1, item.started_at or datetime.min))


async def _infer_workflow_client_id(
    workflow: Workflow,
    db: AsyncSession,
    org_id: str,
) -> str | None:
    agent_ids: list[str] = []
    for node in (workflow.nodes or []) or []:
        data = (node or {}).get("data", {}) or {}
        agent_id = data.get("agent_id")
        if isinstance(agent_id, str) and agent_id:
            agent_ids.append(agent_id)
    if not agent_ids:
        return None

    rows = (
        await db.execute(
            select(Agent.client_id).where(
                Agent.org_id == org_id,
                Agent.id.in_(agent_ids),
                Agent.client_id.is_not(None),
            )
        )
    ).all()
    client_ids = {client_id for (client_id,) in rows if client_id}
    if len(client_ids) == 1:
        return next(iter(client_ids))
    return None


def _serialize_execution_steps(
    execution: Execution,
    *,
    ordered_agents: list[tuple[str, str]],
    agent_names_by_id: dict[str, str],
) -> list[dict[str, Any]]:
    current_index = 0
    serialized: list[dict[str, Any]] = []

    for step in sorted((execution.steps or []), key=lambda item: item.step_index):
        agent_id: str | None = None
        agent_name: str | None = None

        if step.step_type == "update":
            agent_name = "Standup Summary"
        elif current_index < len(ordered_agents):
            agent_id = ordered_agents[current_index][0]
            agent_name = agent_names_by_id.get(agent_id) or ordered_agents[current_index][1]

        serialized.append(
            {
                "id": step.id,
                "execution_id": step.execution_id,
                "org_id": step.org_id,
                "step_type": step.step_type,
                "content": step.content,
                "tool_name": step.tool_name,
                "tool_input": step.tool_input,
                "tool_output": step.tool_output,
                "tool_success": step.tool_success,
                "step_index": step.step_index,
                "duration_ms": step.duration_ms,
                "tokens_used": step.tokens_used,
                "created_at": step.created_at,
                "timestamp": step.created_at,
                "agent_id": agent_id,
                "agent_name": agent_name,
            }
        )

        if step.step_type in {"final_answer", "error"} and current_index < len(ordered_agents):
            current_index += 1

    return serialized


async def run_workflow_background(
    execution_id: str,
    workflow_id: str,
    input_message: str,
    user_id: str | None = None,
    org_id: str | None = None,
    memory_service=None,
    hitl_service=None,
):
    async with AsyncSessionLocal() as db:
        try:
            wf_result = await db.execute(select(Workflow).where(Workflow.id == workflow_id, Workflow.org_id == org_id))
            workflow = wf_result.scalar_one_or_none()
            if not workflow:
                async with AsyncSessionLocal() as fail_db:
                    fail_exec = await fail_db.scalar(
                        select(Execution).where(Execution.id == execution_id)
                    )
                    if fail_exec and fail_exec.status in {ExecutionStatus.pending, ExecutionStatus.running}:
                        fail_exec.status = ExecutionStatus.failed
                        fail_exec.error = "Workflow not found or not accessible"
                        fail_exec.completed_at = datetime.utcnow()
                        await fail_db.commit()
                return

            agent_ids = set()
            for node in workflow.nodes or []:
                data = node.get("data", {}) or {}
                if data.get("agent_id"):
                    agent_ids.add(data["agent_id"])
                for agent_id in data.get("agent_ids") or node.get("agent_ids") or []:
                    if agent_id:
                        agent_ids.add(agent_id)
            agents_result = await db.execute(select(Agent).where(Agent.id.in_(list(agent_ids)), Agent.org_id == org_id))
            agents = agents_result.scalars().all()
            unassigned = [
                n.get("data", {}).get("label", n.get("id"))
                for n in (workflow.nodes or [])
                if not n.get("data", {}).get("agent_id")
                and n.get("type") != "condition"
                and n.get("type") != "parallel_group"
                and n.get("type") != "approval"
                and not n.get("hitl_enabled")
                and not (n.get("config", {}) or {}).get("hitl_enabled")
                and not (n.get("data", {}) or {}).get("hitl_enabled")
                and not ((n.get("data", {}) or {}).get("config", {}) or {}).get("hitl_enabled")
            ]

            await ws_manager.broadcast_to_channel(
                f"org:{org_id}",
                {
                    "type": "execution_start",
                    "execution_id": execution_id,
                    "workflow": workflow.name,
                    "input": input_message,
                    "node_count": len(workflow.nodes or []),
                    "agent_count": len(agents),
                    "unassigned_nodes": unassigned,
                },
            )

            engine = WorkflowEngine(
                db,
                memory_service=memory_service,
                hitl_service=hitl_service,
            )
            output, tokens = await engine.run(workflow_id, input_message, user_id, execution_id)
            execution = await db.scalar(select(Execution).where(Execution.id == execution_id, Execution.org_id == org_id))

            if execution and execution.status == ExecutionStatus.pending_review:
                return

            await ws_manager.broadcast_to_channel(
                f"org:{org_id}",
                {
                    "type": "execution_complete",
                    "execution_id": execution_id,
                    "output": output[:500],
                    "tokens": tokens,
                    "cost": execution.cost if execution else 0.0,
                },
            )
            await ws_manager.broadcast_to_channel(
                f"execution:{execution_id}",
                {
                    "event": "execution_complete",
                    "execution_id": execution_id,
                    "status": "completed",
                    "result_preview": output[:200],
                    "tokens": tokens,
                    "cost": execution.cost if execution else 0.0,
                },
            )

        except WorkflowExecutionStopped as e:
            logger.info("Workflow execution %s stopped with status %s", execution_id, e.status)
            stopped_status = str(e.status.value if hasattr(e.status, "value") else e.status)
            await ws_manager.broadcast_to_channel(
                f"execution:{execution_id}",
                {
                    "event": "execution_complete",
                    "execution_id": execution_id,
                    "status": stopped_status,
                    "result_preview": e.output[:200] if getattr(e, "output", None) else "",
                },
            )
            await ws_manager.broadcast_to_channel(
                f"org:{org_id}",
                {
                    "type": "execution_complete",
                    "execution_id": execution_id,
                    "status": stopped_status,
                    "output": e.output[:500] if getattr(e, "output", None) else "",
                },
            )

        except Exception as e:
            terminal_status = None
            try:
                async with AsyncSessionLocal() as fail_db:
                    fail_exec = await fail_db.scalar(
                        select(Execution).where(Execution.id == execution_id)
                    )
                    terminal_status = (
                        fail_exec.status.value
                        if fail_exec and hasattr(fail_exec.status, "value")
                        else fail_exec.status
                        if fail_exec
                        else None
                    )
                    if fail_exec and fail_exec.status in {ExecutionStatus.pending, ExecutionStatus.running}:
                        fail_exec.status = ExecutionStatus.failed
                        fail_exec.error = str(e)[:500]
                        fail_exec.completed_at = datetime.utcnow()
                        await fail_db.commit()
                        terminal_status = ExecutionStatus.failed.value
            except Exception as db_err:
                logger.error(f"Failed to update execution status on error: {db_err}")

            await ws_manager.broadcast_to_channel(
                f"org:{org_id}",
                {
                    "type": "execution_error",
                    "execution_id": execution_id,
                    "status": terminal_status or ExecutionStatus.failed.value,
                    "error": str(e),
                },
            )
            await ws_manager.broadcast_to_channel(
                f"execution:{execution_id}",
                {
                    "event": "execution_failed",
                    "execution_id": execution_id,
                    "status": terminal_status or ExecutionStatus.failed.value,
                    "error": str(e)[:500],
                },
            )


async def enqueue_workflow_execution(
    execution_id: str,
    workflow_id: str,
    input_message: str,
    user_id: str | None,
    org_id: str | None,
    background_tasks: BackgroundTasks | None = None,
    memory_service=None,
    hitl_service=None,
) -> str:
    from config import settings

    celery_available = _is_celery_broker_reachable(settings.celery_broker_url)
    if celery_available:
        try:
            from tasks.workflow_tasks import run_workflow_task

            run_workflow_task.delay(workflow_id, input_message, user_id or "system", execution_id)
            return "celery"
        except Exception as exc:
            logger.warning("Celery workflow dispatch failed, falling back to local background task: %s", exc)

    if background_tasks is not None:
        background_tasks.add_task(
            run_workflow_background,
            execution_id,
            workflow_id,
            input_message,
            user_id,
            org_id,
            memory_service,
            hitl_service,
        )
        return "background"

    raise RuntimeError("Workflow dispatch unavailable")


def _coerce_workflow_input_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _build_workflow_input_message(workflow: Workflow, data: WorkflowRunRequest) -> str:
    definitions = workflow.input_variables or []
    if not definitions:
        return data.input_message.strip()

    parts: list[str] = []
    missing_labels: list[str] = []

    for raw_var in definitions:
        if not isinstance(raw_var, dict):
            continue
        name = str(raw_var.get("name") or "").strip()
        if not name:
            continue

        label = str(raw_var.get("label") or name).strip() or name
        default_value = _coerce_workflow_input_value(raw_var.get("default", ""))
        value = _coerce_workflow_input_value(data.input_values.get(name, default_value))
        required = bool(raw_var.get("required"))

        if required and not value:
            missing_labels.append(label)
            continue
        if value:
            parts.append(f"{label}: {value}")

    if missing_labels:
        raise HTTPException(
            status_code=422,
            detail=f"Missing required workflow inputs: {', '.join(missing_labels)}",
        )

    if parts:
        return "\n".join(parts)
    return data.input_message.strip()


async def _resolve_execution_client_id(
    *,
    workflow: Workflow,
    data: WorkflowRunRequest,
    db: AsyncSession,
    org_id: str,
) -> str | None:
    if data.client_id:
        client = await db.scalar(
            select(Client).where(
                Client.id == data.client_id,
                Client.org_id == org_id,
            )
        )
        if not client:
            raise HTTPException(status_code=404, detail="Client not found")
        return str(client.id)
    return await _infer_workflow_client_id(workflow, db, org_id)


@router.post("/workflows/{workflow_id}/run", response_model=ExecutionRunResponse, status_code=202)
@limiter.limit("10/minute")
async def run_workflow(
    request: Request,
    response: Response,
    workflow_id: str,
    data: WorkflowRunRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    wf_result = await db.execute(select(Workflow).where(Workflow.id == workflow_id, Workflow.org_id == ctx.org.id))
    workflow = wf_result.scalar_one_or_none()
    if not workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    await check_plan_limit("executions", ctx.org, db)
    input_message = _build_workflow_input_message(workflow, data)
    client_id = await _resolve_execution_client_id(
        workflow=workflow,
        data=data,
        db=db,
        org_id=ctx.org.id,
    )
    execution = Execution(
        id=str(uuid4()),
        org_id=ctx.org.id,
        workflow_id=workflow_id,
        client_id=client_id,
        trigger=data.trigger,
        status=ExecutionStatus.pending,
        input_message=input_message,
        started_at=datetime.utcnow(),
        max_runtime_seconds=data.max_runtime_seconds or 3600,
    )
    db.add(execution)
    await db.commit()
    await db.refresh(execution)

    await enqueue_workflow_execution(
        execution.id,
        workflow_id,
        input_message,
        current_user.id,
        ctx.org.id,
        background_tasks=background_tasks,
        memory_service=getattr(request.app.state, "memory_service", None),
        hitl_service=getattr(request.app.state, "hitl_service", None),
    )

    return ExecutionRunResponse(
        id=execution.id,
        execution_id=execution.id,
        status="queued",
        websocket_channel=f"execution:{execution.id}",
        message="Execution started. Connect to WebSocket for live updates.",
    )


@router.get("", response_model=List[ExecutionResponse])
async def list_executions(
    workflow_id: Optional[str] = None,
    limit: int = 50,
    db: AsyncSession = Depends(get_db),
    ctx: OrgContext = Depends(get_org_context),
):
    query = select(Execution).where(Execution.org_id == ctx.org.id).order_by(Execution.started_at.desc()).limit(limit)
    if workflow_id:
        query = query.where(Execution.workflow_id == workflow_id)
    result = await db.execute(query)
    return result.scalars().all()


@router.get("/{execution_id}", response_model=ExecutionDetailResponse)
async def get_execution(execution_id: str, db: AsyncSession = Depends(get_db), ctx: OrgContext = Depends(get_org_context)):
    result = await db.execute(
        select(Execution)
        .options(selectinload(Execution.steps), selectinload(Execution.workflow))
        .where(Execution.id == execution_id, Execution.org_id == ctx.org.id)
    )
    execution = result.scalar_one_or_none()
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")

    workflow_name = execution.workflow.name if execution.workflow else None
    client_name = None
    agent_name = "Agent"
    model_name = "Default"
    primary_agent_id = _primary_workflow_agent_id(execution.workflow)

    if execution.client_id:
        client = await db.scalar(
            select(Client).where(Client.id == execution.client_id, Client.org_id == ctx.org.id)
        )
        if client:
            client_name = client.company_name or client.name

    if primary_agent_id:
        agent_result = await db.execute(
            select(Agent).where(Agent.id == primary_agent_id, Agent.org_id == ctx.org.id)
        )
        agent = agent_result.scalar_one_or_none()
        if agent:
            agent_name = agent.name or agent_name
            if agent.model_config_id:
                from database.models import ModelConfig

                config_result = await db.execute(
                    select(ModelConfig.display_name).where(
                        ModelConfig.id == agent.model_config_id,
                        ModelConfig.org_id == ctx.org.id,
                    )
                )
                model_name = config_result.scalar_one_or_none() or agent.model or model_name
            else:
                model_name = agent.model or model_name

    ordered_agents = _ordered_workflow_agents(execution.workflow)
    agent_name_map: dict[str, str] = {}
    if ordered_agents:
        agent_rows = await db.execute(
            select(Agent.id, Agent.name).where(
                Agent.id.in_([agent_id for agent_id, _ in ordered_agents]),
                Agent.org_id == ctx.org.id,
            )
        )
        agent_name_map = {agent_id: agent_name for agent_id, agent_name in agent_rows.all()}

    return {
        "id": execution.id,
        "workflow_id": execution.workflow_id,
        "client_id": execution.client_id,
        "client_name": client_name,
        "parent_execution_id": execution.parent_execution_id,
        "workflow_name": workflow_name,
        "agent_name": agent_name,
        "model_name": model_name,
        "trigger": execution.trigger,
        "status": execution.status.value if hasattr(execution.status, "value") else str(execution.status),
        "input": execution.input_message,
        "input_message": execution.input_message,
        "output": _execution_output_text(execution),
        "output_message": execution.output_message,
        "revision_number": execution.revision_number,
        "ceo_feedback": execution.ceo_feedback,
        "started_at": execution.started_at,
        "completed_at": execution.completed_at,
        "approved_by": execution.approved_by,
        "approved_at": execution.approved_at,
        "approval_note": execution.approval_note,
        "delivered_at": execution.delivered_at,
        "delivery_method": execution.delivery_method,
        "delivery_target": execution.delivery_target,
        "token_count": execution.token_count,
        "cost": execution.cost,
        "error": execution.error,
        "warning": execution.warning,
        "max_runtime_seconds": execution.max_runtime_seconds,
        "steps": _serialize_execution_steps(
            execution,
            ordered_agents=ordered_agents,
            agent_names_by_id=agent_name_map,
        ),
    }


@router.get("/{execution_id}/export")
async def export_execution(
    execution_id: str,
    format: str = Query(..., pattern="^(pdf|docx)$"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    result = await db.execute(
        select(Execution)
        .options(selectinload(Execution.steps), selectinload(Execution.workflow))
        .where(Execution.id == execution_id, Execution.org_id == ctx.org.id)
    )
    execution = result.scalar_one_or_none()
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")

    output = _execution_output_text(execution).strip()
    if execution.status != ExecutionStatus.completed or not execution.approved_at or not output:
        raise HTTPException(status_code=400, detail="Execution must be approved before export")

    workflow_name = execution.workflow.name if execution.workflow else "Execution Report"
    primary_agent_id = _primary_workflow_agent_id(execution.workflow)
    agent_name = "Agent"
    if primary_agent_id:
        agent = await db.scalar(
            select(Agent).where(Agent.id == primary_agent_id, Agent.org_id == ctx.org.id)
        )
        if agent:
            agent_name = agent.persona_name or agent.name or agent_name

    client_name: str | None = None
    if execution.client_id:
        client = await db.scalar(
            select(Client).where(Client.id == execution.client_id, Client.org_id == ctx.org.id)
        )
        if client:
            client_name = client.company_name or client.name or None

    approved_by_name = current_user.full_name or current_user.email
    if execution.approved_by:
        approver = await db.scalar(select(User).where(User.id == execution.approved_by))
        if approver:
            approved_by_name = approver.full_name or approver.email or approved_by_name

    exporter = _render_execution_pdf if format == "pdf" else _render_execution_docx
    try:
        content = await exporter(
            agency_name=ctx.org.name,
            workflow_name=workflow_name,
            client_name=client_name,
            completed_at=execution.completed_at,
            agent_name=agent_name,
            approved_by_name=approved_by_name,
            approved_at=execution.approved_at,
            output=output,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    media_type = (
        "application/pdf"
        if format == "pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    filename = _build_export_filename(workflow_name, execution.completed_at, format)
    return StreamingResponse(
        BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/{execution_id}/deliver", response_model=ExecutionDeliverResponse, status_code=200)
async def deliver_execution(
    execution_id: str,
    data: ExecutionDeliverRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    result = await db.execute(
        select(Execution)
        .options(selectinload(Execution.steps), selectinload(Execution.workflow))
        .where(Execution.id == execution_id, Execution.org_id == ctx.org.id)
    )
    execution = result.scalar_one_or_none()
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")

    output = _execution_output_text(execution).strip()
    if execution.status != ExecutionStatus.completed or not execution.approved_at or not output:
        raise HTTPException(status_code=400, detail="Execution must be approved before delivery")

    workflow_name = execution.workflow.name if execution.workflow else "Execution"
    client: Client | None = None
    client_name: str | None = None
    if execution.client_id:
        client = await db.scalar(
            select(Client).where(Client.id == execution.client_id, Client.org_id == ctx.org.id)
        )
        if client:
            client_name = client.company_name or client.name or None

    delivered_at = datetime.utcnow()
    target: str

    if data.method == "email":
        email_to = (data.email_to or "").strip()
        if not email_to:
            raise HTTPException(status_code=422, detail="email_to is required for email delivery")
        integration = await db.scalar(
            select(UserIntegration).where(
                UserIntegration.org_id == ctx.org.id,
                UserIntegration.user_id == current_user.id,
                UserIntegration.integration_type == IntegrationType.gmail,
                UserIntegration.is_active == True,  # noqa: E712
            )
        )
        if not integration:
            raise HTTPException(
                status_code=422,
                detail="Gmail not connected. Connect it in /integrations",
            )
        gmail_config = decrypt_config(integration.config)

        client_line = (
            f'<p style="color:#475569;">Client: {html.escape(client_name)}</p>'
            if client_name
            else ""
        )
        html_body = (
            '<html><body style="font-family: Arial, sans-serif; line-height: 1.6; color: #0f172a;">'
            f'<h1 style="font-size: 22px; margin-bottom: 8px;">{html.escape(workflow_name)}</h1>'
            f"{client_line}"
            f"{_markdown_to_html(output)}"
            "</body></html>"
        )
        text_body = _markdown_to_plain_text(output)

        pdf_bytes: bytes | None = None
        primary_agent_id = _primary_workflow_agent_id(execution.workflow)
        agent_name = "Agent"
        if primary_agent_id:
            agent = await db.scalar(
                select(Agent).where(Agent.id == primary_agent_id, Agent.org_id == ctx.org.id)
            )
            if agent:
                agent_name = agent.persona_name or agent.name or agent_name
        try:
            pdf_bytes = await _render_execution_pdf(
                agency_name=ctx.org.name,
                workflow_name=workflow_name,
                client_name=client_name,
                completed_at=execution.completed_at,
                agent_name=agent_name,
                approved_by_name=current_user.full_name or current_user.email,
                approved_at=execution.approved_at,
                output=output,
            )
        except Exception as exc:
            logger.warning("Failed to attach PDF for execution delivery %s: %s", execution.id, exc)

        try:
            await _send_execution_email_via_gmail(
                org_id=ctx.org.id,
                user_id=current_user.id,
                to_email=email_to,
                subject=_delivery_subject(workflow_name, client_name),
                html_body=html_body,
                text_body=text_body,
                pdf_bytes=pdf_bytes,
                pdf_filename=_build_export_filename(workflow_name, execution.completed_at, "pdf"),
                prefetched_config=gmail_config,
            )
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        target = email_to

    elif data.method == "google_doc":
        integration = await db.scalar(
            select(UserIntegration).where(
                UserIntegration.org_id == ctx.org.id,
                UserIntegration.user_id == current_user.id,
                UserIntegration.integration_type == IntegrationType.gmail,
                UserIntegration.is_active == True,  # noqa: E712
            )
        )
        if not integration:
            raise HTTPException(
                status_code=422,
                detail="Google not connected. Connect it in /integrations",
            )
        gmail_config = decrypt_config(integration.config)
        title = (data.doc_title or "").strip() or (
            f"{workflow_name} — {client_name or 'Client'} — {_format_export_date(execution.completed_at)}"
        )
        try:
            target = await _create_google_doc(
                org_id=ctx.org.id,
                user_id=current_user.id,
                title=title,
                content=output,
                prefetched_config=gmail_config,
            )
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc

    else:
        if not client:
            raise HTTPException(status_code=422, detail="Assign this execution to a client first")
        if not client.portal_token:
            client.portal_token = secrets.token_urlsafe(32)
        if not client.portal_enabled:
            client.portal_enabled = True
            client.updated_at = delivered_at
        target = f"{str(request.base_url).rstrip('/')}/portal/{client.portal_token}"

    execution.delivered_at = delivered_at
    execution.delivery_method = data.method
    execution.delivery_target = target
    await db.commit()
    await db.refresh(execution)

    return ExecutionDeliverResponse(
        delivered=True,
        method=data.method,
        target=target,
        delivered_at=execution.delivered_at,
    )


@router.get("/{execution_id}/revisions", response_model=List[ExecutionRevisionResponse])
async def get_execution_revisions(
    execution_id: str,
    db: AsyncSession = Depends(get_db),
    ctx: OrgContext = Depends(get_org_context),
):
    execution = await db.scalar(
        select(Execution).where(Execution.id == execution_id, Execution.org_id == ctx.org.id)
    )
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")

    revisions = await _load_execution_revision_chain(execution, db, ctx.org.id)
    return [
        {
            "id": item.id,
            "revision_number": item.revision_number or 1,
            "status": item.status.value if hasattr(item.status, "value") else str(item.status),
            "ceo_feedback": item.ceo_feedback,
            "output": _execution_output_text(item),
            "started_at": item.started_at,
            "approved_at": item.approved_at,
            "approved_by": item.approved_by,
        }
        for item in revisions
    ]


@router.post("/{execution_id}/approve", status_code=200)
async def approve_execution(
    execution_id: str,
    data: ExecutionApproveRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    result = await db.execute(
        select(Execution)
        .options(selectinload(Execution.workflow), selectinload(Execution.steps))
        .where(Execution.id == execution_id, Execution.org_id == ctx.org.id)
    )
    execution = result.scalar_one_or_none()
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")
    if execution.status != ExecutionStatus.pending_review:
        raise HTTPException(status_code=400, detail="Execution is not pending review")

    execution.status = ExecutionStatus.completed
    execution.approved_by = str(current_user.id)
    execution.approved_at = datetime.utcnow()
    if data.note:
        execution.approval_note = data.note
    await db.commit()

    primary_agent_id = _primary_workflow_agent_id(execution.workflow)
    if primary_agent_id:
        from services.trust_score_service import trust_score_service

        await trust_score_service.record_task_completed(
            agent_id=primary_agent_id,
            success=True,
            on_time=True,
            cost_cents=int((execution.cost or 0) * 100),
            budget_cents=10000,
            tools_used=[step.tool_name for step in (execution.steps or []) if step.tool_name],
            db=db,
        )
        await trust_score_service.record_review_result(
            agent_id=primary_agent_id,
            passed=True,
            db=db,
        )

    if data.note and primary_agent_id:
        preference_text = await _extract_preferences(data.note, primary_agent_id)
        await _store_ceo_preference(
            agent_id=primary_agent_id,
            org_id=str(execution.org_id),
            preference_text=preference_text,
            db=db,
            source="ceo_feedback",
        )

    from services.cto_memory_service import cto_memory_service

    workflow_name = execution.workflow.name if execution.workflow else "workflow"
    approval_context = f"{workflow_name}: {(execution.input_message or '')[:120]}"
    await cto_memory_service.record_approval_pattern(
        org_id=ctx.org.id,
        action_type="execution_review",
        context=approval_context,
        was_approved=True,
        db=db,
    )

    await ws_manager.broadcast_to_channel(
        f"org:{ctx.org.id}",
        {
            "event": "execution_complete",
            "execution_id": execution_id,
            "status": ExecutionStatus.completed.value,
            "output": (execution.output_message or "")[:500],
        },
    )
    await ws_manager.broadcast_to_channel(
        f"execution:{execution_id}",
        {
            "event": "execution_complete",
            "execution_id": execution_id,
            "status": ExecutionStatus.completed.value,
            "result_preview": (execution.output_message or "")[:200],
        },
    )

    return {"status": "completed", "execution_id": execution_id}


@router.post("/{execution_id}/regenerate", response_model=ExecutionRegenerateResponse, status_code=202)
async def regenerate_execution(
    execution_id: str,
    data: ExecutionRegenerateRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    result = await db.execute(
        select(Execution)
        .options(selectinload(Execution.workflow), selectinload(Execution.steps))
        .where(Execution.id == execution_id, Execution.org_id == ctx.org.id)
    )
    source_execution = result.scalar_one_or_none()
    if not source_execution:
        raise HTTPException(status_code=404, detail="Execution not found")
    feedback = data.feedback.strip()
    if not feedback:
        raise HTTPException(status_code=422, detail="Feedback is required to regenerate")
    if source_execution.status not in {
        ExecutionStatus.pending_review,
        ExecutionStatus.completed,
        ExecutionStatus.rejected,
        ExecutionStatus.failed,
    }:
        raise HTTPException(status_code=400, detail="Execution cannot be regenerated in its current state")
    if not source_execution.workflow:
        raise HTTPException(status_code=404, detail="Workflow not found")

    primary_agent_id = _primary_workflow_agent_id(source_execution.workflow)
    if feedback and primary_agent_id:
        from services.trust_score_service import trust_score_service

        await trust_score_service.record_review_result(
            agent_id=primary_agent_id,
            passed=False,
            db=db,
        )
        preference_text = await _extract_preferences(feedback, primary_agent_id)
        await _store_ceo_preference(
            agent_id=primary_agent_id,
            org_id=ctx.org.id,
            preference_text=preference_text,
            db=db,
            source="ceo_feedback",
        )

    revised_input = _build_regeneration_prompt(source_execution, feedback)
    revision_id = str(uuid4())
    execution = Execution(
        id=revision_id,
        org_id=ctx.org.id,
        workflow_id=source_execution.workflow_id,
        client_id=source_execution.client_id,
        parent_execution_id=source_execution.id,
        trigger="revision",
        status=ExecutionStatus.pending,
        input_message=revised_input,
        revision_number=(source_execution.revision_number or 1) + 1,
        ceo_feedback=feedback,
        started_at=datetime.utcnow(),
        max_runtime_seconds=source_execution.max_runtime_seconds or 3600,
    )
    db.add(execution)
    source_execution.status = ExecutionStatus.cancelled
    source_execution.error = f"Superseded by revision {revision_id}"
    source_execution.completed_at = source_execution.completed_at or datetime.utcnow()
    await db.commit()
    await db.refresh(execution)

    await enqueue_workflow_execution(
        revision_id,
        source_execution.workflow_id,
        revised_input,
        current_user.id,
        ctx.org.id,
        background_tasks=background_tasks,
        memory_service=getattr(background_tasks, "memory_service", None),
        hitl_service=getattr(background_tasks, "hitl_service", None),
    )

    return ExecutionRegenerateResponse(
        revision_id=revision_id,
        revision_number=execution.revision_number,
        status="running",
    )


@router.get("/{execution_id}/messages", response_model=List[MessageResponse])
async def get_execution_messages(
    execution_id: str,
    db: AsyncSession = Depends(get_db),
    ctx: OrgContext = Depends(get_org_context),
):
    execution = await db.scalar(select(Execution).where(Execution.id == execution_id, Execution.org_id == ctx.org.id))
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")
    result = await db.execute(
        select(Message).where(Message.execution_id == execution_id).order_by(Message.timestamp)
    )
    return result.scalars().all()


@router.delete("/{execution_id}", status_code=200)
async def cancel_execution(
    execution_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    execution = await db.scalar(
        select(Execution)
        .where(Execution.id == execution_id, Execution.org_id == ctx.org.id)
    )
    if not execution:
        raise HTTPException(status_code=404, detail="Execution not found")
    if execution.status not in (ExecutionStatus.running, ExecutionStatus.pending):
        raise HTTPException(
            status_code=400,
            detail=f"Cannot cancel execution with status '{execution.status.value}'"
        )
    execution.status = ExecutionStatus.cancelled
    execution.error = f"Cancelled by {current_user.email}"
    execution.completed_at = datetime.utcnow()
    await db.commit()

    await ws_manager.broadcast_to_channel(
        f"execution:{execution_id}",
        {
            "event": "execution_failed",
            "execution_id": execution_id,
            "error": "Cancelled by user",
            "status": "cancelled",
        },
    )
    return {"cancelled": True, "execution_id": execution_id}
