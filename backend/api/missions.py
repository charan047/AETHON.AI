from __future__ import annotations

from datetime import datetime
from io import BytesIO
import html
import secrets

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from auth.dependencies import get_current_user, require_editor
from auth.org_context import OrgContext, get_org_context
from database import get_db
from database.models import Client, Execution, ExecutionStatus, Mission, MissionStatus, MissionTask, MissionTaskStatus, User
from services.goal_decomposer import goal_decomposer
from tasks.mission_tasks import run_mission_task


router = APIRouter(dependencies=[Depends(get_current_user), Depends(get_org_context)])


class MissionCreate(BaseModel):
    goal: str = Field(..., min_length=1)
    client_id: str | None = None


class MissionTaskResponse(BaseModel):
    id: str
    mission_id: str
    org_id: str
    sequence: int
    title: str
    description: str | None = None
    agent_id: str | None = None
    depends_on: str | None = None
    status: str
    output_summary: str | None = None
    output_file_id: str | None = None
    execution_id: str | None = None
    started_at: datetime | None = None
    completed_at: datetime | None = None

    model_config = {"from_attributes": True}


class MissionStatsResponse(BaseModel):
    total: int
    pending: int
    running: int
    completed: int
    failed: int
    skipped: int


class MissionResponse(BaseModel):
    id: str
    org_id: str
    client_id: str | None = None
    client_name: str | None = None
    goal: str
    title: str | None = None
    status: str
    report: str | None = None
    report_delivered: bool
    created_by: str | None = None
    created_at: datetime
    completed_at: datetime | None = None
    stats: MissionStatsResponse
    tasks: list[MissionTaskResponse]


class MissionReportResponse(BaseModel):
    mission_id: str
    status: str
    report: str | None = None


def _build_mission_export_filename(title: str | None, completed_at: datetime | None, ext: str) -> str:
    safe_title = (title or "mission-report").strip().lower().replace(" ", "-")
    safe_title = "".join(ch for ch in safe_title if ch.isalnum() or ch in {"-", "_"}).strip("-_") or "mission-report"
    date_part = (completed_at or datetime.utcnow()).strftime("%Y-%m-%d")
    return f"{safe_title}-{date_part}.{ext}"


def _parse_markdown_blocks(content: str) -> list[dict]:
    blocks: list[dict] = []
    lines = (content or "").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            index += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue
        if stripped.startswith("##"):
            level = len(stripped) - len(stripped.lstrip("#"))
            blocks.append({"type": "heading", "level": level, "text": stripped[level:].strip()})
            index += 1
            continue
        if stripped.startswith(("- ", "* ")):
            items: list[str] = []
            while index < len(lines):
                current = lines[index].strip()
                if not current.startswith(("- ", "* ")):
                    break
                items.append(current[2:].strip())
                index += 1
            blocks.append({"type": "bullet", "items": items})
            continue
        if stripped[:2].isdigit() and ". " in stripped:
            items: list[str] = []
            while index < len(lines):
                current = lines[index].strip()
                if ". " not in current:
                    break
                prefix, body = current.split(". ", 1)
                if not prefix.isdigit():
                    break
                items.append(body.strip())
                index += 1
            blocks.append({"type": "numbered", "items": items})
            continue
        paragraph: list[str] = [stripped]
        index += 1
        while index < len(lines):
            current = lines[index].strip()
            if not current or current.startswith(("##", "- ", "* ", "```")):
                break
            paragraph.append(current)
            index += 1
        blocks.append({"type": "paragraph", "text": "\n".join(paragraph)})
    return blocks


async def _render_mission_report_pdf(
    *,
    agency_name: str,
    mission_title: str,
    client_name: str | None,
    completed_at: datetime | None,
    report: str,
    tasks: list[MissionTask],
) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:
        raise RuntimeError("PDF export dependency is not installed") from exc

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=1.0 * inch,
        bottomMargin=0.85 * inch,
        title=mission_title,
        author=agency_name,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="MissionTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=20, leading=24, textColor=colors.HexColor("#111827"), spaceAfter=14))
    styles.add(ParagraphStyle(name="MissionMeta", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13, textColor=colors.HexColor("#6B7280"), spaceAfter=4))
    styles.add(ParagraphStyle(name="MissionH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=colors.HexColor("#111827"), spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="MissionH3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=12, leading=16, textColor=colors.HexColor("#1F2937"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="MissionBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=16, textColor=colors.HexColor("#111827"), spaceAfter=8))
    styles.add(ParagraphStyle(name="MissionCode", parent=styles["Code"], fontName="Courier", fontSize=9, leading=12, textColor=colors.HexColor("#111827"), backColor=colors.HexColor("#F8FAFC"), borderColor=colors.HexColor("#CBD5E1"), borderWidth=0.5, borderPadding=8, spaceAfter=10))

    story = [
        Paragraph(mission_title + (f" — {client_name}" if client_name else ""), styles["MissionTitle"]),
        Paragraph((completed_at or datetime.utcnow()).strftime("%b %d, %Y %I:%M %p"), styles["MissionMeta"]),
        Paragraph(f"Agency: {agency_name}", styles["MissionMeta"]),
    ]
    if client_name:
        story.append(Paragraph(f"Client: {client_name}", styles["MissionMeta"]))
    story.append(Spacer(1, 10))

    for block in _parse_markdown_blocks(report):
        if block["type"] == "heading":
            style_name = "MissionH2" if block["level"] <= 2 else "MissionH3"
            story.append(Paragraph(html.escape(block["text"]), styles[style_name]))
        elif block["type"] == "bullet":
            items = [ListItem(Paragraph(html.escape(item), styles["MissionBody"])) for item in block["items"]]
            story.append(ListFlowable(items, bulletType="bullet", leftIndent=18))
            story.append(Spacer(1, 8))
        elif block["type"] == "numbered":
            items = [ListItem(Paragraph(html.escape(item), styles["MissionBody"])) for item in block["items"]]
            story.append(ListFlowable(items, bulletType="1", leftIndent=18))
            story.append(Spacer(1, 8))
        elif block["type"] == "code":
            story.append(Paragraph(html.escape(block["text"]).replace("\n", "<br/>"), styles["MissionCode"]))
        else:
            story.append(Paragraph(html.escape(block["text"]).replace("\n", "<br/>"), styles["MissionBody"]))

    if tasks:
        story.append(Spacer(1, 10))
        story.append(Paragraph("Task Breakdown", styles["MissionH2"]))
        for task in tasks:
            title = f"{task.sequence}. {task.title} — {_status_value(task.status)}"
            story.append(Paragraph(html.escape(title), styles["MissionH3"]))
            if task.output_summary:
                story.append(Paragraph(html.escape(task.output_summary).replace("\n", "<br/>"), styles["MissionBody"]))

    doc.build(story)
    return buffer.getvalue()


async def _render_mission_report_docx(
    *,
    agency_name: str,
    mission_title: str,
    client_name: str | None,
    completed_at: datetime | None,
    report: str,
    tasks: list[MissionTask],
) -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise RuntimeError("DOCX export dependency is not installed") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.add_run(mission_title + (f" — {client_name}" if client_name else "")).bold = True

    for text in (
        (completed_at or datetime.utcnow()).strftime("%b %d, %Y %I:%M %p"),
        f"Agency: {agency_name}",
        f"Client: {client_name}" if client_name else None,
    ):
        if not text:
            continue
        document.add_paragraph(str(text))

    for block in _parse_markdown_blocks(report):
        if block["type"] == "heading":
            level = 2 if block["level"] <= 2 else 3
            document.add_heading(block["text"], level=level)
        elif block["type"] == "bullet":
            for item in block["items"]:
                document.add_paragraph(item, style="List Bullet")
        elif block["type"] == "numbered":
            for item in block["items"]:
                document.add_paragraph(item, style="List Number")
        elif block["type"] == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            document.add_paragraph(block["text"])

    if tasks:
        document.add_heading("Task Breakdown", level=2)
        for task in tasks:
            document.add_heading(f"{task.sequence}. {task.title} — {_status_value(task.status)}", level=3)
            if task.output_summary:
                document.add_paragraph(task.output_summary)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _status_value(value) -> str:
    return value.value if hasattr(value, "value") else str(value)


def _task_stats(tasks: list[MissionTask]) -> MissionStatsResponse:
    counts = {
        "total": len(tasks),
        "pending": 0,
        "running": 0,
        "completed": 0,
        "failed": 0,
        "skipped": 0,
    }
    for task in tasks:
        key = _status_value(task.status)
        if key in counts:
            counts[key] += 1
    return MissionStatsResponse(**counts)


def _serialize_task(task: MissionTask) -> MissionTaskResponse:
    return MissionTaskResponse(
        id=task.id,
        mission_id=task.mission_id,
        org_id=task.org_id,
        sequence=task.sequence,
        title=task.title,
        description=task.description,
        agent_id=task.agent_id,
        depends_on=task.depends_on,
        status=_status_value(task.status),
        output_summary=task.output_summary,
        execution_id=task.execution_id,
        started_at=task.started_at,
        completed_at=task.completed_at,
    )


def _serialize_mission(
    mission: Mission,
    tasks: list[MissionTask],
    client_name: str | None = None,
) -> MissionResponse:
    return MissionResponse(
        id=mission.id,
        org_id=mission.org_id,
        client_id=mission.client_id,
        client_name=client_name,
        goal=mission.goal,
        title=mission.title,
        status=_status_value(mission.status),
        report=mission.report,
        report_delivered=mission.report_delivered,
        created_by=mission.created_by,
        created_at=mission.created_at,
        completed_at=mission.completed_at,
        stats=_task_stats(tasks),
        tasks=[_serialize_task(task) for task in tasks],
    )


async def _get_org_scoped_mission(mission_id: str, org_id: str, db: AsyncSession) -> Mission:
    mission = await db.scalar(
        select(Mission).where(Mission.id == mission_id, Mission.org_id == org_id)
    )
    if not mission:
        raise HTTPException(status_code=404, detail="Mission not found")
    return mission


async def _mission_tasks(mission_id: str, org_id: str, db: AsyncSession) -> list[MissionTask]:
    result = await db.execute(
        select(MissionTask)
        .where(MissionTask.mission_id == mission_id, MissionTask.org_id == org_id)
        .order_by(MissionTask.sequence.asc())
    )
    return list(result.scalars().all())


@router.post("/missions", response_model=MissionResponse, status_code=status.HTTP_201_CREATED)
async def create_mission(
    payload: MissionCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    goal = payload.goal.strip()
    if not goal:
        raise HTTPException(status_code=422, detail="Goal is required")

    client_name = None
    if payload.client_id:
        client = await db.scalar(
            select(Client).where(Client.id == payload.client_id, Client.org_id == ctx.org.id)
        )
        if not client:
            raise HTTPException(status_code=404, detail="Client not found")
        client_name = client.company_name or client.name

    mission = await goal_decomposer.create_mission(
        goal=goal,
        org_id=ctx.org.id,
        client_id=payload.client_id,
        created_by=current_user.id,
        db=db,
    )
    tasks = await _mission_tasks(mission.id, ctx.org.id, db)
    run_mission_task.delay(str(mission.id))
    return _serialize_mission(mission, tasks, client_name=client_name)


@router.get("/missions", response_model=list[MissionResponse])
async def list_missions(
    db: AsyncSession = Depends(get_db),
    ctx: OrgContext = Depends(get_org_context),
):
    missions = (
        await db.execute(
            select(Mission)
            .where(Mission.org_id == ctx.org.id)
            .order_by(Mission.created_at.desc())
        )
    ).scalars().all()

    clients = (
        await db.execute(
            select(Client).where(Client.org_id == ctx.org.id)
        )
    ).scalars().all()
    client_names = {client.id: (client.company_name or client.name) for client in clients}

    payload: list[MissionResponse] = []
    for mission in missions:
        tasks = await _mission_tasks(mission.id, ctx.org.id, db)
        payload.append(
            _serialize_mission(
                mission,
                tasks,
                client_name=client_names.get(mission.client_id),
            )
        )
    return payload


@router.get("/missions/{mission_id}", response_model=MissionResponse)
async def get_mission(
    mission_id: str,
    db: AsyncSession = Depends(get_db),
    ctx: OrgContext = Depends(get_org_context),
):
    mission = await _get_org_scoped_mission(mission_id, ctx.org.id, db)
    tasks = await _mission_tasks(mission.id, ctx.org.id, db)
    client_name = None
    if mission.client_id:
        client = await db.scalar(
            select(Client).where(Client.id == mission.client_id, Client.org_id == ctx.org.id)
        )
        if client:
            client_name = client.company_name or client.name
    return _serialize_mission(mission, tasks, client_name=client_name)


@router.get("/missions/{mission_id}/report", response_model=MissionReportResponse)
async def get_mission_report(
    mission_id: str,
    db: AsyncSession = Depends(get_db),
    ctx: OrgContext = Depends(get_org_context),
):
    mission = await _get_org_scoped_mission(mission_id, ctx.org.id, db)
    return MissionReportResponse(
        mission_id=mission.id,
        status=_status_value(mission.status),
        report=mission.report,
    )


@router.get("/missions/{mission_id}/export")
async def export_mission_report(
    mission_id: str,
    format: str = Query(..., pattern="^(pdf|docx)$"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    mission = await _get_org_scoped_mission(mission_id, ctx.org.id, db)
    if not mission.report:
        raise HTTPException(status_code=400, detail="Mission report is not ready yet")

    tasks = await _mission_tasks(mission.id, ctx.org.id, db)
    client_name = None
    if mission.client_id:
        client = await db.scalar(
            select(Client).where(Client.id == mission.client_id, Client.org_id == ctx.org.id)
        )
        if client:
            client_name = client.company_name or client.name

    exporter = _render_mission_report_pdf if format == "pdf" else _render_mission_report_docx
    try:
        content = await exporter(
            agency_name=ctx.org.name,
            mission_title=mission.title or mission.goal or "Mission Report",
            client_name=client_name,
            completed_at=mission.completed_at,
            report=mission.report,
            tasks=tasks,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    media_type = (
        "application/pdf"
        if format == "pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    filename = _build_mission_export_filename(mission.title or mission.goal, mission.completed_at, format)
    return StreamingResponse(
        BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/missions/{mission_id}/approve-report", response_model=MissionResponse, status_code=200)
async def approve_mission_report(
    mission_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    mission = await _get_org_scoped_mission(mission_id, ctx.org.id, db)
    if _status_value(mission.status) != MissionStatus.completed.value:
        raise HTTPException(status_code=400, detail="Mission must be completed before approval")
    if not mission.report:
        raise HTTPException(status_code=400, detail="Mission report is not ready")
    if not mission.client_id:
        raise HTTPException(status_code=400, detail="Mission is not linked to a client")

    client = await db.scalar(
        select(Client).where(Client.id == mission.client_id, Client.org_id == ctx.org.id)
    )
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")

    approved_at = datetime.utcnow()
    if not client.portal_token:
        client.portal_token = secrets.token_urlsafe(32)
    if not client.portal_enabled:
        client.portal_enabled = True
    client.updated_at = approved_at

    mission.report_delivered = True
    await db.commit()

    from services.cto_memory_service import cto_memory_service

    client_label = client.company_name or client.name or "client"
    await cto_memory_service.record_approval_pattern(
        org_id=ctx.org.id,
        action_type="deliver_portal",
        context=f"{mission.title or mission.goal} for {client_label}",
        was_approved=True,
        db=db,
    )

    tasks = await _mission_tasks(mission.id, ctx.org.id, db)
    return _serialize_mission(
        mission,
        tasks,
        client_name=client.company_name or client.name,
    )


@router.post("/missions/{mission_id}/retry", response_model=MissionResponse, status_code=status.HTTP_201_CREATED)
async def retry_mission(
    mission_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    original = await _get_org_scoped_mission(mission_id, ctx.org.id, db)

    client_name = None
    if original.client_id:
        client = await db.scalar(
            select(Client).where(Client.id == original.client_id, Client.org_id == ctx.org.id)
        )
        if client:
            client_name = client.company_name or client.name

    mission = await goal_decomposer.create_mission(
        goal=original.goal,
        org_id=ctx.org.id,
        client_id=original.client_id,
        created_by=current_user.id,
        db=db,
    )
    tasks = await _mission_tasks(mission.id, ctx.org.id, db)
    run_mission_task.delay(str(mission.id))
    return _serialize_mission(mission, tasks, client_name=client_name)


@router.delete("/missions/{mission_id}")
async def delete_mission(
    mission_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_editor),
    ctx: OrgContext = Depends(get_org_context),
):
    mission = await _get_org_scoped_mission(mission_id, ctx.org.id, db)
    tasks = await _mission_tasks(mission.id, ctx.org.id, db)

    execution_ids = [task.execution_id for task in tasks if task.execution_id]
    if execution_ids:
        executions = (
            await db.execute(
                select(Execution).where(
                    Execution.id.in_(execution_ids),
                    Execution.org_id == ctx.org.id,
                )
            )
        ).scalars().all()
        for execution in executions:
            if execution.status not in {
                ExecutionStatus.completed,
                ExecutionStatus.failed,
                ExecutionStatus.cancelled,
                ExecutionStatus.timed_out,
            }:
                execution.status = ExecutionStatus.cancelled
                execution.error = "Mission cancelled by user"
                execution.completed_at = datetime.utcnow()

    for task in tasks:
        if task.status in {MissionTaskStatus.pending, MissionTaskStatus.running}:
            task.status = MissionTaskStatus.failed if task.execution_id else MissionTaskStatus.skipped
            task.completed_at = datetime.utcnow()
            task.output_summary = task.output_summary or "Mission cancelled by user"

    mission.status = MissionStatus.failed
    mission.completed_at = datetime.utcnow()
    await db.commit()
    return {"success": True, "mission_id": mission.id, "status": "failed"}
